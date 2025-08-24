from PyQt5.QtWidgets import (
    QWidget, QVBoxLayout, QHBoxLayout, QTextEdit, QLineEdit, QPushButton,
    QListWidget, QLabel, QFileDialog
)
from PyQt5.QtGui import QFont, QTextCursor
from PyQt5.QtCore import Qt, QTimer, QPropertyAnimation, QEasingCurve
from core.doc_loader import extract_text_from_docx, extract_text_from_files, extract_text_from_pdf
from core.ai_engine import ask_question_with_context
from core.vector_store import VectorSearch
from ui.settings_window import SettingsWindow
from core.settings_manager import load_settings
import markdown2
from os.path import basename
from docx import Document
from docx.shared import Pt
from PyQt5.QtWidgets import QFileDialog
import os
from ui.settings_panel import SettingsPanel


class MainWindow(QWidget):
    def __init__(self):
        super().__init__()
        self.uploaded_files = []
        self.document_text = ""
        self.vector_search = VectorSearch()
        self.settings = load_settings()
        self.model_name = self.settings["model_name"]
        self.chunk_size = self.settings["chunk_size"]
        self.llm_source = self.settings.get("llm_source", "ollama")  # default to 'ollama' or whatever you use
        self.dark_mode = self.settings["dark_mode"]

        self.setWindowTitle("LexiAI – Your AI Document Assistant")
        self.setGeometry(200, 200, 1000, 700)

        self.init_ui()
        self.init_settings_panel()
        self.apply_theme()

    def init_ui(self):
        main_layout = QHBoxLayout()

        # LEFT PANEL
        left_panel = QVBoxLayout()

        title = QLabel("LexiAI")
        title.setFont(QFont("Segoe UI", 36, QFont.Bold))
        title.setAlignment(Qt.AlignmentFlag.AlignCenter)
        left_panel.addWidget(title)

        self.response_area = QTextEdit()
        self.response_area.setReadOnly(True)
        self.response_area.setFont(QFont("Segoe UI", 26))
        self.response_area.setStyleSheet("""
            QTextEdit {
                font-size: 26px;
                padding: 16px;
                border-radius: 18px;
                background-color: rgba(255, 255, 255, 0.05);
                color: #f1f5f9;
                border: 1px solid rgba(255, 255, 255, 0.1);
                
            }
        """)
        left_panel.addWidget(self.response_area, stretch=1)

        self.typing_label = QLabel("")
        self.typing_label.setFont(QFont("Segoe UI", 18, QFont.StyleItalic))
        self.typing_label.setStyleSheet("color: #9ca3af; padding-left: 10px;")
        left_panel.addWidget(self.typing_label)

        bottom_layout = QHBoxLayout()

        self.question_input = QLineEdit()
        self.question_input.setFont(QFont("Segoe UI", 26))
        self.question_input.setPlaceholderText("Ask LexiAI something...")
        self.question_input.setStyleSheet("""
            QLineEdit {
                font-size: 26px;
                padding: 12px;
                border-radius: 8px;
                background-color: #2a2a3d;
                color: white;
            }
        """)
        self.question_input.returnPressed.connect(self.process_question)
        bottom_layout.addWidget(self.question_input, stretch=1)

        self.upload_button = QPushButton("+")
        self.upload_button.setFixedSize(50, 50)
        self.upload_button.setFont(QFont("Segoe UI", 24, QFont.Bold))
        self.upload_button.setStyleSheet("""
            QPushButton {
                font-size: 24px;
                font-weight: bold;
                border-radius: 25px;
                background-color: #3b82f6;
                color: white;
            }
            QPushButton:hover {
                background-color: #2563eb;
            }
        """)
        self.upload_button.clicked.connect(self.upload_files)
        bottom_layout.addWidget(self.upload_button)

        self.save_button = QPushButton("💾")
        self.save_button.setFixedSize(50, 50)
        self.save_button.setFont(QFont("Segoe UI", 24))
        self.save_button.setStyleSheet("""
            QPushButton {
                font-size: 24px;
                border-radius: 25px;
                background-color: #10b981;
                color: white;
            }
            QPushButton:hover {
                background-color: #059669;
            }
        """)
        self.save_button.clicked.connect(self.save_chat)
        bottom_layout.addWidget(self.save_button)

        self.folder_button = QPushButton("📁")
        self.folder_button.setFixedSize(50, 50)
        self.folder_button.setFont(QFont("Segoe UI", 24, QFont.Bold))
        self.folder_button.setStyleSheet("""
            QPushButton {
                font-size: 24px;
                font-weight: bold;
                border-radius: 25px;
                background-color: #10b981;
                color: white;
            }
            QPushButton:hover {
                background-color: #059669;
            }
        """)
        self.folder_button.clicked.connect(self.upload_folder)
        bottom_layout.addWidget(self.folder_button)



        # Remove (-) Button
        self.remove_button = QPushButton("–")
        self.remove_button.setFixedSize(50, 50)
        self.remove_button.setFont(QFont("Segoe UI", 30, QFont.Bold))
        self.remove_button.setStyleSheet("""
            QPushButton {
                font-size: 30px;
                font-weight: bold;
                border-radius: 25px;
                background-color: #ef4444;
                color: white;
            }
            QPushButton:hover {
                background-color: #dc2626;
            }
        """)
        self.remove_button.clicked.connect(self.remove_selected_file)
        bottom_layout.addWidget(self.remove_button)


        left_panel.addLayout(bottom_layout)
        left_widget = QWidget()
        left_widget.setLayout(left_panel)

        # RIGHT SIDEBAR
        self.sidebar_visible = True

        self.file_list = QListWidget()
        self.file_list.setFont(QFont("Segoe UI", 28))
        self.file_list.setStyleSheet("""
            QListWidget {
                font-size: 28px;
                background-color: #1a1a2b;
                color: #e0e0e0;
                border-radius: 12px;
                padding: 10px;
                border: 1px solid #3b82f6;
            }
        """)
      

        self.settings_button = QPushButton("⚙")
        self.settings_button.setFixedSize(40, 40)
        self.settings_button.setFont(QFont("Segoe UI", 26))
        self.settings_button.clicked.connect(self.open_settings)
        self.settings_button.setStyleSheet("""
            QPushButton {
                background-color: transparent;
                color: white;
                font-size: 26px;
            }
            QPushButton:hover {
                color: #3b82f6;
            }
        """)

        label = QLabel("📄 Files")
        label.setFont(QFont("Segoe UI", 24))

        sidebar_header = QHBoxLayout()
        sidebar_header.addWidget(label)
        sidebar_header.addStretch()
        sidebar_header.addWidget(self.settings_button)

        sidebar_top = QWidget()
        sidebar_top.setLayout(sidebar_header)

        sidebar_layout = QVBoxLayout()
        sidebar_layout.addWidget(sidebar_top)
        # Manual Search UI
        self.search_input = QLineEdit()
        self.search_input.setPlaceholderText("Search in uploaded files...")
        self.search_input.setFont(QFont("Segoe UI", 18))
        self.search_input.setStyleSheet("padding: 6px; background-color: #2a2a3d; color: white;")
        self.search_input.setStyleSheet("""
            QLineEdit {
                background-color: #12121c;
                color: #ffffff;
                padding: 8px 12px;
                border: 1px solid #3b82f6;
                border-radius: 8px;
                font-size: 18px;
            }
        """)


        self.search_button = QPushButton("🔍")
        self.search_button.setFixedHeight(40)
        self.search_button.setFont(QFont("Segoe UI", 18))
        self.search_button.setStyleSheet("""
            QPushButton {
                background-color: #3b82f6;
                color: white;
                border-radius: 10px;
                padding: 8px 14px;
            }
            QPushButton:hover {
                background-color: #2563eb;
            }
        """)

        self.search_button.clicked.connect(self.search_in_files)

        search_bar_layout = QHBoxLayout()
        search_bar_layout.addWidget(self.search_input)
        search_bar_layout.addWidget(self.search_button)

        search_bar_widget = QWidget()
        search_bar_widget.setLayout(search_bar_layout)
        sidebar_layout.addWidget(search_bar_widget)

        sidebar_layout.addWidget(self.file_list)

        self.sidebar_widget = QWidget()
        self.sidebar_widget.setLayout(sidebar_layout)
        self.sidebar_widget.setMinimumWidth(180)
        self.sidebar_widget.setMaximumWidth(240)
        self.sidebar_widget.setStyleSheet("""
            QWidget {
                background-color: #111827;
                border-left: 1px solid rgba(59, 130, 246, 0.2);
                border-radius: 0px;
            }
            QListWidget {
                background-color: #1f2937;
                border: none;
                border-radius: 8px;
                padding: 8px;
                color: #e5e7eb;
                font-size: 20px;
            }
            QLineEdit {
                background-color: #1f2937;
                border: 1px solid #3b82f6;
                border-radius: 8px;
                padding: 6px 10px;
                font-size: 16px;
                color: #e5e7eb;
            }
        """)


        self.toggle_sidebar_btn = QPushButton("⏵")
        self.toggle_sidebar_btn.setFixedSize(40, 100)
        self.toggle_sidebar_btn.setFont(QFont("Segoe UI", 20))
        self.toggle_sidebar_btn.clicked.connect(self.toggle_sidebar)
        
        
        self.toggle_sidebar_btn.setStyleSheet("""
            QPushButton {
                background-color: #1f2937;
                color: #3b82f6;
                border: 1px solid rgba(59, 130, 246, 0.3);
                border-top-left-radius: 10px;
                border-bottom-left-radius: 10px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #2563eb;
                color: white;
            }
        """)
        self.toggle_sidebar_btn.setFixedWidth(20)
        self.toggle_sidebar_btn.setFixedHeight(80)

       # WRAPPER for sidebar and toggle button
        self.sidebar_wrapper = QWidget()
        self.sidebar_layout = QHBoxLayout()
        self.sidebar_layout.setContentsMargins(0, 0, 0, 0)
        self.sidebar_layout.setSpacing(0)

        self.sidebar_layout.addWidget(self.sidebar_widget)
        self.sidebar_layout.addWidget(self.toggle_sidebar_btn)

        self.sidebar_wrapper.setLayout(self.sidebar_layout)
        self.sidebar_layout.addStretch()
        

        # Set fixed width for right panel
        #right_widget.setFixedWidth(self.sidebar_widget.width() + self.toggle_sidebar_btn.width() + 1)

        # Combine panels
        main_layout.addWidget(left_widget, stretch=1)
        main_layout.addWidget(self.sidebar_wrapper, stretch=0)
        self.setLayout(main_layout)
        self.main_layout = main_layout  # Store reference for later use

    def toggle_sidebar(self):
        if self.sidebar_visible:
            self.animation = QPropertyAnimation(self.sidebar_wrapper, b"maximumWidth")
            self.animation.setDuration(300)
            self.animation.setStartValue(self.sidebar_wrapper.width())
            self.animation.setEndValue(self.toggle_sidebar_btn.width())  # Only show the button
            self.animation.setEasingCurve(QEasingCurve.InOutCubic)
            self.animation.start()

            self.sidebar_widget.setVisible(False)
            self.sidebar_visible = False
            self.toggle_sidebar_btn.setText("⏴")
        else:
            self.sidebar_widget.setVisible(True)

            self.animation = QPropertyAnimation(self.sidebar_wrapper, b"maximumWidth")
            self.animation.setDuration(300)
            self.animation.setStartValue(self.sidebar_wrapper.width())
            self.animation.setEndValue(280)  # sidebar 240 + button 40
            self.animation.setEasingCurve(QEasingCurve.InOutCubic)
            self.animation.start()

            self.sidebar_visible = True
            self.toggle_sidebar_btn.setText("⏵")



    def start_typing_animation(self):
        self.typing_dots = 0
        self.typing_label.setText("LexiAI is typing")
        self.dot_timer = QTimer()
        self.dot_timer.timeout.connect(self.update_typing_label)
        self.dot_timer.start(500)

    def update_typing_label(self):
        self.typing_dots = (self.typing_dots + 1) % 4
        self.typing_label.setText("LexiAI is typing" + "." * self.typing_dots)

    def stop_typing_animation(self):
        if hasattr(self, 'dot_timer'):
            self.dot_timer.stop()
        self.typing_label.setText("")

    def process_question(self):
        question = self.question_input.text().strip()
        if not question:
            self.response_area.append("⚠️ Please enter a question.")
            return

        self.question_input.clear()
        self.response_area.moveCursor(QTextCursor.End)
        self.response_area.insertHtml(f"<b>You:</b> {question}<br><br>")
        self.response_area.insertHtml(f"<b>LexiAI:</b> ")
        
        # Get relevant document chunks only if files are uploaded
        context = ""
        if self.document_text.strip():
            relevant_chunks = self.vector_search.query(question)
            context = "\n".join(relevant_chunks)

        # Determine assistant role from settings
        ROLE_MAP = {
            "General": "general",
            "Legal Advisor": "legal",
            "HR Assistant": "hr",
            "Technical Expert": "tech"
        }
        role_setting = self.settings.get("role", "General")
        role = ROLE_MAP.get(role_setting, "general")
        # Ask question with or without context
        print(f"[DEBUG] Role setting from config: {role_setting}")
        print(f"[DEBUG] Mapped internal role key: {role}")
        answer = ask_question_with_context(
            question,
            context,
            model_name=self.model_name,
            role=role
        )

        self.typing_text = answer
        self.typing_index = 0
        self.start_typing_animation()
        self.typing_cursor = self.response_area.textCursor()
        self._type_text_markdown()

    def init_settings_panel(self):
        self.settings_panel = SettingsPanel(self)
        self.settings_panel.setVisible(False)
        self.main_layout.addWidget(self.settings_panel)

    def search_in_files(self):
        keyword = self.search_input.text().strip().lower()
        if not keyword:
            self.response_area.append("⚠️ Enter a keyword to search.")
            return

        results = []
        for display_name, file_path in self.file_path_map.items():
            content = ""
            if file_path.endswith(".pdf"):
                content = extract_text_from_pdf(file_path).lower()
            elif file_path.endswith(".docx"):
                content = extract_text_from_docx(file_path).lower()

            if keyword in content:
                match_preview = content[max(0, content.find(keyword) - 100): content.find(keyword) + 100]
                results.append(f"<b>{display_name}</b><br>{match_preview}<br><br>")

        if results:
            self.response_area.append("<b>🔎 Search Results:</b><br>" + "".join(results))
        else:
            self.response_area.append("❌ No matches found.")


    def save_chat(self):
        options = QFileDialog.Options()
        file_path, _ = QFileDialog.getSaveFileName(
            self,
            "Save Chat",
            "chat_history",
            "Text Files (*.txt);;HTML Files (*.html);;Word Files (*.docx)",
            options=options
        )

        if file_path:
            if file_path.endswith(".txt"):
                content = self.response_area.toPlainText()
                with open(file_path, "w", encoding="utf-8") as file:
                    file.write(content)

            elif file_path.endswith(".html"):
                content = self.response_area.toHtml()
                with open(file_path, "w", encoding="utf-8") as file:
                    file.write(content)

            elif file_path.endswith(".docx"):
                doc = Document()
                
                lines = self.response_area.toPlainText().split("\n")
                for line in lines:
                    if line.strip():
                        p = doc.add_paragraph(line.strip())
                        # Set font for the paragraph's run
                        if p.runs:
                            p.runs[0].font.name = 'Segoe UI'
                            p.runs[0].font.size = Pt(12)
                        else:
                            # If no run exists, create one
                            run = p.add_run(line.strip())
                            run.font.name = 'Segoe UI'
                            run.font.size = Pt(12)

                doc.save(file_path)
    def _type_text_markdown(self):
        if not self.typing_text:
            self.stop_typing_animation()
            return

        def type_char():
            if self.typing_index < len(self.typing_text):
                self.typing_cursor.insertText(self.typing_text[self.typing_index])
                self.typing_index += 1
            else:
                self.typing_timer.stop()
                self.stop_typing_animation()
                self.typing_cursor.insertHtml("<br><br>")

        self.typing_timer = QTimer()
        self.typing_timer.timeout.connect(type_char)
        self.typing_timer.start(15)


    def upload_files(self):
        # Let user choose files *or* a folder
        file_dialog = QFileDialog(self)
        file_dialog.setFileMode(QFileDialog.ExistingFiles)
        file_dialog.setNameFilters(["Documents (*.pdf *.docx)"])
        file_dialog.setOption(QFileDialog.DontUseNativeDialog, True)

        # Add a folder selection button to dialog
        folder_button = file_dialog.findChild(QPushButton, "lookInCombo")
        if folder_button:
            folder_button.setText("Browse Folder")

        if file_dialog.exec_():
            selected_files = file_dialog.selectedFiles()
            all_files = []

            for path in selected_files:
                if os.path.isdir(path):
                    for root, _, files in os.walk(path):
                        for file in files:
                            if file.endswith((".pdf", ".docx")):
                                full_path = os.path.join(root, file)
                                all_files.append(full_path)
                else:
                    all_files.append(path)

            new_files = [f for f in all_files if f not in self.uploaded_files]
            if new_files:
                self.uploaded_files.extend(new_files)
                self.file_path_map = getattr(self, 'file_path_map', {})
                for f in new_files:
                    display_name = os.path.basename(f)
                    self.file_list.addItem(display_name)
                    self.file_path_map[display_name] = f

                new_text = extract_text_from_files(new_files)
                self.document_text += "\n" + new_text
                self.vector_search.add_documents(new_text)
                print(f"[Loaded Files] {self.uploaded_files}")
                print(f"[Total Combined Characters] {len(self.document_text)}")

                preview = self.document_text[:1000]
                print(f"[Combined Preview]\n{preview}...")

    

    def upload_folder(self):
        folder_path = QFileDialog.getExistingDirectory(self, "Select Folder")
        if folder_path:
            all_files = []
            for root, _, files in os.walk(folder_path):
                for file in files:
                    if file.endswith((".pdf", ".docx")):
                        full_path = os.path.join(root, file)
                        all_files.append(full_path)

            new_files = [f for f in all_files if f not in self.uploaded_files]
            if new_files:
                self.uploaded_files.extend(new_files)
                self.file_path_map = getattr(self, 'file_path_map', {})
                for f in new_files:
                    display_name = os.path.basename(f)
                    self.file_list.addItem(display_name)
                    self.file_path_map[display_name] = f

                
                new_text = extract_text_from_files(new_files)
                self.document_text += "\n" + new_text
                self.vector_search.add_documents(new_text)
                print(f"[Loaded Files] {self.uploaded_files}")
                print(f"[Total Combined Characters] {len(self.document_text)}")

                preview = self.document_text[:1000]
                print(f"[Combined Preview]\n{preview}...")



    def remove_selected_file(self):
        selected_items = self.file_list.selectedItems()
        if not selected_items:
            return

        for item in selected_items:
            file_name = item.text()
            file_path = self.file_path_map.get(file_name)
            if file_path and file_path in self.uploaded_files:
                self.uploaded_files.remove(file_path)
                del self.file_path_map[file_name]

            else:
                print(f"[⚠] File not found in uploaded_files: {file_name}")

            row = self.file_list.row(item)
            self.file_list.takeItem(row)

        # Rebuild document context if files remain
        if self.uploaded_files:
            self.document_text = extract_text_from_files(self.uploaded_files)
            self.vector_search = VectorSearch()
            self.vector_search.add_documents(self.document_text)
        else:
            self.document_text = ""
            self.vector_search = VectorSearch()
        print("[Uploaded Files]", self.uploaded_files)


    def open_settings(self):
      self.settings_panel.setVisible(not self.settings_panel.isVisible())

    def apply_theme(self):
        if self.dark_mode:
            self.setStyleSheet(self.load_styles())
        else:
            self.setStyleSheet("""
            QWidget {
                background-color: #ffffff;
                color: #000000;
                font-family: 'Segoe UI', sans-serif;
            }
            """)

    def apply_updated_settings(self):
        self.settings = load_settings()
        self.model_name = self.settings.get("model_name", "llama3")
        self.chunk_size = self.settings.get("chunk_size", 500)
        self.llm_source = self.settings.get("llm_source", "ollama")
        self.dark_mode = self.settings.get("dark_mode", True)

        self.apply_theme()
        print("🔄 Settings reloaded and UI updated.")


    def load_styles(self):
        return """
        QWidget {
            background-color: #1e1e2f;
            color: #ffffff;
            font-family: 'Segoe UI', sans-serif;
            font-size: 24px;
        }
        QPushButton {
            background-color: #3b82f6;
            color: white;
            border: none;
            border-radius: 8px;
        }
        QPushButton:hover {
            background-color: #2563eb;
        }
        QListWidget {
            background-color: #2a2a3d;
            border-radius: 8px;
            padding: 6px;
        }
        QLabel {
            color: #e0e0e0;
        }
        QScrollBar:vertical, QScrollBar:horizontal {
            background: #2a2a3d;
            border-radius: 7px;
        }
        QScrollBar::handle:vertical, QScrollBar::handle:horizontal {
            background: #3b82f6;
            min-height: 20px;
            border-radius: 7px;
        }
        QScrollBar::add-line, QScrollBar::sub-line {
            height: 0px; width: 0px;
        }
        """
